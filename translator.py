from docx import Document
from googletrans import Translator
import logging

# Set up logging
logging.basicConfig(filename='translation_errors.log', level=logging.ERROR)

def translate_document(input_file, output_file, src_lang='en', dest_lang='kn', chunk_size=5000):
    translator = Translator()
    doc = Document(input_file)
    translated_doc = Document()

    for paragraph in doc.paragraphs:
        text = paragraph.text
        try:
            # Process in chunks to handle large documents
            chunks = [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]
            translated_text = ''
            for chunk in chunks:
                translation = translator.translate(chunk, src=src_lang, dest=dest_lang)
                if translation and translation.text:
                    translated_text += translation.text
                else:
                    translated_text += "[Translation failed]"
                    logging.error(f"Translation failed for chunk: {chunk}")
        except Exception as e:
            translated_text = f"[Translation error: {str(e)}]"
            logging.error(f"Translation error for paragraph: {text}\nError: {str(e)}")
        
        translated_doc.add_paragraph(translated_text)
    
    translated_doc.save(output_file)

if __name__ == "__main__":
    input_file = 'input.docx'
    output_file = 'output.docx'
    translate_document(input_file, output_file)
